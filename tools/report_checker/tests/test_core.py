"""
查重分析 — 核心逻辑测试
"""
import sys, os, struct, zlib, tempfile
from io import BytesIO
sys.stdout.reconfigure(encoding='utf-8')
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', '..', '..'))

from tools.report_checker.core import (
    extract_features, generate_excel, analyze,
    unzip_process, organize_process,
)

PASS = FAIL = 0


def check(cond, msg):
    global PASS, FAIL
    if cond: PASS += 1; print(f'  ✓ {msg}')
    else: FAIL += 1; print(f'  ✗ {msg}')


def check_eq(a, b, msg):
    global PASS, FAIL
    if a == b: PASS += 1; print(f'  ✓ {msg} ({a})')
    else: FAIL += 1; print(f'  ✗ {msg}: got {a}, expected {b}')


def section(name):
    print(f'\n{"="*60}\n  {name}\n{"="*60}')


def _make_png(seed=0, w=4, h=4):
    """Create a minimal valid PNG with seed-varied content."""
    r = (seed * 50 + 100) % 256
    g = (seed * 80 + 50) % 256
    b_val = (seed * 120 + 200) % 256
    raw = b''
    for y in range(h):
        raw += b'\x00' + bytes([r, g, b_val]) * w
    def chunk(ctype, data):
        c = ctype + data
        return struct.pack('>I', len(data)) + c + struct.pack('>I', zlib.crc32(c) & 0xffffffff)
    ihdr = struct.pack('>IIBBBBB', w, h, 8, 2, 0, 0, 0)
    return (
        b'\x89PNG\r\n\x1a\n'
        + chunk(b'IHDR', ihdr)
        + chunk(b'IDAT', zlib.compress(raw))
        + chunk(b'IEND', b'')
    )


def _make_docx(path, png_count=1):
    """Create a docx file with specified number of embedded images."""
    from docx import Document
    from docx.shared import Inches
    doc = Document()
    for i in range(png_count):
        doc.add_picture(BytesIO(_make_png(seed=i)), width=Inches(1))
    doc.save(path)


# ════════════════════════════════════════════
section('Test 1: extract_features')
# ════════════════════════════════════════════

with tempfile.TemporaryDirectory() as td:
    p = os.path.join(td, "test.docx")
    _make_docx(p, 2)
    feats = extract_features(p)
    check_eq(len(feats), 2, '提取到 2 张图片指纹')
    check(isinstance(feats[0][0], str) and len(feats[0][0]) == 32,
          '指纹为 MD5 十六进制字符串')
    check(isinstance(feats[0][1], int) and feats[0][1] > 0,
          '字节大小为正整数')

    # 无图片文档
    from docx import Document
    doc = Document()
    doc.save(os.path.join(td, "empty.docx"))
    feats = extract_features(os.path.join(td, "empty.docx"))
    check_eq(len(feats), 0, '无图片文档返回空列表')

    # 无效文件
    feats = extract_features(os.path.join(td, "nonexist.docx"))
    check_eq(len(feats), 0, '不存在文件返回空列表')


# ════════════════════════════════════════════
section('Test 2: extract_features 相同图片返回相同指纹')
# ════════════════════════════════════════════

with tempfile.TemporaryDirectory() as td:
    _make_docx(os.path.join(td, "a.docx"), 1)
    _make_docx(os.path.join(td, "b.docx"), 1)
    fa = extract_features(os.path.join(td, "a.docx"))
    fb = extract_features(os.path.join(td, "b.docx"))
    check_eq(fa, fb, '相同图片生成相同指纹')

    # 不同图片数量
    _make_docx(os.path.join(td, "c.docx"), 2)
    fc = extract_features(os.path.join(td, "c.docx"))
    check_eq(len(fc), 2, '不同图片数量指纹不同')


# ════════════════════════════════════════════
section('Test 3: analyze 交叉比对 — 比例 < 40% 判定正常')
# ════════════════════════════════════════════

with tempfile.TemporaryDirectory() as td:
    #   张三: 6 张图片 (seed 0~5)
    #   李四: 6 张图片 (seed 0~5，与张三全相同) → 6/6=100% ≥ 40% → 抄袭
    #   王五: 20 张图片 (seed 0~19，前 6 张与张三相同) → 6/20=30% < 40% → 正常
    _make_docx(os.path.join(td, "04230001_张三.docx"), 6)
    _make_docx(os.path.join(td, "04230002_李四.docx"), 6)
    _make_docx(os.path.join(td, "04230003_王五.docx"), 20)

    r = analyze(td)
    check_eq(r['total'], 3, '共 3 份报告')
    check_eq(r['plag_count'], 2, '2 份抄袭（张三、李四 100%；王五 30% < 40% 正常）')
    for row in r['results']:
        name = row[0]
        if '王五' in name:
            check(row[1] == '正常', f'{name} 判定正常（6/20=30% < 40%）')


# ════════════════════════════════════════════
section('Test 4: analyze 判定抄袭')
# ════════════════════════════════════════════

with tempfile.TemporaryDirectory() as td:
    _make_docx(os.path.join(td, "04230001_张三.docx"), 5)
    _make_docx(os.path.join(td, "04230002_李四.docx"), 5)
    # 王五 1 张独立图片 (seed=99，与其他文件不重复)
    from docx import Document
    from docx.shared import Inches
    doc = Document()
    doc.add_picture(BytesIO(_make_png(seed=99)), width=Inches(1))
    doc.save(os.path.join(td, "04230003_王五.docx"))

    r = analyze(td)
    check_eq(r['total'], 3, '共 3 份报告')
    check(r['plag_count'] >= 2, f'至少 2 份抄袭（实际 {r["plag_count"]} 份）')

    for row in r['results']:
        if '张三' in row[0]:
            check(row[1] == '疑似抄袭', f'{row[0]} 判定抄袭（5/5=100%）')
        if '李四' in row[0]:
            check(row[1] == '疑似抄袭', f'{row[0]} 判定抄袭（5/5=100%）')
        if '王五' in row[0]:
            check(row[1] == '正常', f'{row[0]} 判定正常（无重复图片）')


# ════════════════════════════════════════════
section('Test 5: analyze 无图片报告')
# ════════════════════════════════════════════

with tempfile.TemporaryDirectory() as td:
    from docx import Document
    for name in ["04230001_张三.docx", "04230002_李四.docx"]:
        Document().save(os.path.join(td, name))

    r = analyze(td)
    check_eq(r['none_count'], 2, '2 份无图片报告')
    for row in r['results']:
        check(row[1] == '无图片', f'{row[0]} 判定无图片')


# ════════════════════════════════════════════
section('Test 6: generate_excel')
# ════════════════════════════════════════════

buf = generate_excel(
    [["04230001_张三", "疑似抄袭", "与 04230002_李四 有 5 张一致"],
     ["04230002_李四", "正常", "包含 3 张唯一截图"]],
    total=3, plag_count=1, none_count=1,
)
from openpyxl import load_workbook
buf.seek(0)
wb = load_workbook(buf)
sheets = wb.sheetnames
check_eq(len(sheets), 2, 'Excel 包含 2 个工作表')
check('统计概览' in sheets, '包含统计概览表')
check('详细名单' in sheets, '包含详细名单表')
ws2 = wb['详细名单']
rows = list(ws2.iter_rows(values_only=True))
check_eq(len(rows), 3, '详细名单有 2 行数据 + 1 行标题')
check_eq(rows[1][0], '04230001_张三', '第一行数据为张三')
check_eq(rows[1][1], '疑似抄袭', '张三判定为抄袭')
wb.close()


# ════════════════════════════════════════════
section('Test 7: analyze 空目录')
# ════════════════════════════════════════════

with tempfile.TemporaryDirectory() as td:
    r = analyze(td)
    check_eq(r['total'], 0, '空目录返回 total=0')
    check_eq(len(r['results']), 0, '空目录 results 为空')


# ════════════════════════════════════════════
section('Test 8: unzip + organize 流程')
# ════════════════════════════════════════════

with tempfile.TemporaryDirectory() as td:
    # 构造模拟的学生目录结构
    raw_dir = os.path.join(td, "raw")
    os.makedirs(os.path.join(raw_dir, "04230001_张三"))
    os.makedirs(os.path.join(raw_dir, "04230002_李四"))
    os.makedirs(os.path.join(raw_dir, "04230003_王五_其他"))

    _make_docx(os.path.join(raw_dir, "04230001_张三", "实验报告.docx"), 1)
    _make_docx(os.path.join(raw_dir, "04230002_李四", "答题记录.docx"), 1)
    _make_docx(os.path.join(raw_dir, "04230002_李四", "实验报告.docx"), 2)
    _make_docx(os.path.join(raw_dir, "04230003_王五_其他", "报告.docx"), 1)
    # 空目录
    os.makedirs(os.path.join(raw_dir, "04230004_赵六"))

    target_dir = os.path.join(td, "organized")
    organize_process(raw_dir, target_dir)

    files = sorted(os.listdir(target_dir))
    check_eq(len(files), 4, '整理后共 4 个文件（赵六无报告→空白占位，其他 3 人各 1 份）')
    check('04230001_张三.docx' in files, '张三报告正确提取')
    check('04230002_李四.docx' in files, '李四报告正确提取（跳过答题记录）')
    check('王五' in ' '.join(files), '04230003_王五_其他 → 04230003_王五')

    # 提取的图片数量正确
    feats = extract_features(os.path.join(target_dir, "04230002_李四.docx"))
    check_eq(len(feats), 2, '李四报告有 2 张图片')


# ════════════════════════════════════════════
section('Test 9: run_pipeline 端到端')
# ════════════════════════════════════════════

from tools.report_checker.core import run_pipeline

with tempfile.TemporaryDirectory() as td:
    # 构造 ZIP
    student_dirs = {
        "04230001_张三": [("实验报告.docx", 5)],
        "04230002_李四": [("实验报告.docx", 5)],
        "04230003_王五": [("实验报告.docx", 1)],
    }
    raw_dir = os.path.join(td, "src_raw")
    os.makedirs(raw_dir, exist_ok=True)
    for name, files in student_dirs.items():
        d = os.path.join(raw_dir, name)
        os.makedirs(d, exist_ok=True)
        for fname, count in files:
            _make_docx(os.path.join(d, fname), count)

    zip_path = os.path.join(td, "test.zip")
    import zipfile
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as z:
        for root, _, files in os.walk(raw_dir):
            for f in files:
                fpath = os.path.join(root, f)
                arcname = os.path.relpath(fpath, raw_dir)
                z.write(fpath, arcname)

    ws = os.path.join(td, "workspace")
    logs = []
    result = run_pipeline(zip_path, ws, log_func=lambda msg: logs.append(msg))

    check(result['total'] == 3, '共 3 份报告')
    check(result['plag_count'] >= 2, '至少 2 份抄袭')
    check(len(logs) >= 3, f'日志回调被调用至少 3 次（实际 {len(logs)} 次）')
    check('excel_buf' in result, '包含 Excel 缓冲区')
    result['excel_buf'].seek(0)
    check(len(result['excel_buf'].read()) > 0, 'Excel 不为空')


# ════════════════════════════════════════════
section('Test 10: run_pipeline progress_callback=None')
# ════════════════════════════════════════════

with tempfile.TemporaryDirectory() as td:
    student_dirs = {
        "04230001_张三": [("实验报告.docx", 2)],
        "04230002_李四": [("实验报告.docx", 2)],
    }
    raw_dir = os.path.join(td, "src_raw")
    os.makedirs(raw_dir, exist_ok=True)
    for name, files in student_dirs.items():
        d = os.path.join(raw_dir, name)
        os.makedirs(d, exist_ok=True)
        for fname, count in files:
            _make_docx(os.path.join(d, fname), count)

    import zipfile
    zip_path = os.path.join(td, "test.zip")
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as z:
        for root, _, files in os.walk(raw_dir):
            for f in files:
                fpath = os.path.join(root, f)
                arcname = os.path.relpath(fpath, raw_dir)
                z.write(fpath, arcname)

    result = run_pipeline(zip_path, os.path.join(td, "ws"))
    check(result['total'] == 2, '10 无 callback 时 total=2')
    check('excel_buf' in result, '10 无 callback 时有 excel_buf')
    result['excel_buf'].seek(0)
    check(len(result['excel_buf'].read()) > 0, '10 无 callback 时 Excel 不为空')


# ════════════════════════════════════════════
section('Test 11: 阈值边界 — 刚好 40%')
# ════════════════════════════════════════════

with tempfile.TemporaryDirectory() as td:
    # 张三: 5 张 (seed 0~4)
    # 李四: 5 张 (seed 0~4 → 5/5=100% > 40%) → 抄袭
    # 王五: 5 张 (seed 0,1 → 2/5=40% → 刚好等于阈值) → 应判定抄袭
    # 赵六: 5 张 (seed 0 → 1/5=20% < 40%) → 正常
    _make_docx(os.path.join(td, "张三.docx"), 5)
    _make_docx(os.path.join(td, "李四.docx"), 5)
    # 王五的 5 张图：前 2 张与张三共享（seed 0,1），后 3 张独立（seed 90,91,92）
    from docx import Document
    from docx.shared import Inches
    from io import BytesIO
    doc = Document()
    for s in [0, 1, 90, 91, 92]:
        doc.add_picture(BytesIO(_make_png(seed=s)), width=Inches(1))
    doc.save(os.path.join(td, "王五.docx"))
    # 赵六的 5 张图：仅 seed 0 与张三共享
    doc = Document()
    for s in [0, 80, 81, 82, 83]:
        doc.add_picture(BytesIO(_make_png(seed=s)), width=Inches(1))
    doc.save(os.path.join(td, "赵六.docx"))

    r = analyze(td)
    check_eq(r['total'], 4, '11 共 4 份报告')
    for row in r['results']:
        name = row[0]
        if '张三' in name:
            check(row[1] == '疑似抄袭', f'11 {name} 5/5=100% 抄袭')
        elif '李四' in name:
            check(row[1] == '疑似抄袭', f'11 {name} 5/5=100% 抄袭')
        elif '王五' in name:
            check(row[1] == '疑似抄袭', f'11 {name} 2/5=40% 抄袭(≥40%)')
        elif '赵六' in name:
            check(row[1] == '正常', f'11 {name} 1/5=20% 正常(<40%)')


# ════════════════════════════════════════════
section('Test 12: 单一学生 — 无交叉比对')
# ════════════════════════════════════════════

with tempfile.TemporaryDirectory() as td:
    _make_docx(os.path.join(td, "04230001_张三.docx"), 3)
    r = analyze(td)
    check_eq(r['total'], 1, '12 共 1 份报告')
    check_eq(r['plag_count'], 0, '12 单一学生无抄袭')
    check(r['results'][0][1] == '正常', f'12 判定正常: {r["results"][0][1]}')
    check('相同' not in r['results'][0][2], '12 详细依据无相同图片描述')


# ════════════════════════════════════════════
section('Test 13: 完全独立图片 — 全部正常')
# ════════════════════════════════════════════

with tempfile.TemporaryDirectory() as td:
    for i, name in enumerate(["张三", "李四", "王五"]):
        doc = Document()
        doc.add_picture(BytesIO(_make_png(seed=100 + i)), width=Inches(1))
        doc.add_picture(BytesIO(_make_png(seed=200 + i)), width=Inches(1))
        doc.save(os.path.join(td, f"{name}.docx"))

    r = analyze(td)
    check_eq(r['total'], 3, '13 共 3 份报告')
    check_eq(r['plag_count'], 0, '13 无抄袭')
    for row in r['results']:
        check(row[1] == '正常', f'13 {row[0]} 判定正常')


# ════════════════════════════════════════════
section('Test 14: analyze 三条学生数据交集')
# ════════════════════════════════════════════

with tempfile.TemporaryDirectory() as td:
    # 张三: seed 0,1,2
    _make_docx(os.path.join(td, "张三.docx"), 3)
    # 李四: seed 0,1,3 (与张三共享 0,1 → 2/3=66.7% > 40%) → 抄袭
    doc = Document()
    for s in [0, 1, 3]:
        doc.add_picture(BytesIO(_make_png(seed=s)), width=Inches(1))
    doc.save(os.path.join(td, "李四.docx"))
    # 王五: seed 0,4,5 (与张三共享 0 → 1/3=33.3% < 40%) → 正常
    doc = Document()
    for s in [0, 4, 5]:
        doc.add_picture(BytesIO(_make_png(seed=s)), width=Inches(1))
    doc.save(os.path.join(td, "王五.docx"))

    r = analyze(td)
    check_eq(r['total'], 3, '14 共 3 份报告')
    for row in r['results']:
        name = row[0]
        if '张三' in name:
            check(row[1] == '疑似抄袭', f'14 {name} 2/3=66.7%(vs李四) 抄袭')
        elif '李四' in name:
            check(row[1] == '疑似抄袭', f'14 {name} 2/3=66.7% 抄袭')
        elif '王五' in name:
            check(row[1] == '正常', f'14 {name} 1/3=33.3% 正常')
    check_eq(r['plag_count'], 2, '14 2 份抄袭（张三和李四各 2/3=66.7%，王五 1/3=33.3% 正常）')


# ════════════════════════════════════════════
section(f'SUMMARY\n  report_checker: {PASS} 通过, {FAIL} 失败')
if FAIL:
    sys.exit(1)
